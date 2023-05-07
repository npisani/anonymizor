import random
import string
import os
import csv
import re
from docx import Document
import openpyxl
import tkinter as tk
from tkinter import filedialog

def generate_random_string(length):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

def anonymize_run(run, anonymized_terms):
    run_text = run.text
    for term, placeholder in anonymized_terms.items():
        term_pattern = re.compile(re.escape(term), re.IGNORECASE)
        run_text = term_pattern.sub(placeholder, run_text)
    run.text = run_text

def anonymize_docx(input_file, output_file, sensitive_terms):
    document = Document(input_file)
    anonymized_terms = {}

    for term in sensitive_terms:
        placeholder = generate_random_string(len(term))
        anonymized_terms[term] = placeholder

    # Anonymize terms in paragraphs
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            anonymize_run(run, anonymized_terms)

    # Anonymize terms in tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        anonymize_run(run, anonymized_terms)

    document.save(output_file)
    return anonymized_terms

def anonymize_xlsx(input_file, output_file, sensitive_terms):
    workbook = openpyxl.load_workbook(input_file)
    anonymized_terms = {}

    for term in sensitive_terms:
        placeholder = generate_random_string(len(term))
        anonymized_terms[term] = placeholder

        for sheet in workbook:
            for row in sheet.iter_rows():
                for cell in row:
                    if term in str(cell.value):
                        cell.value = str(cell.value).replace(term, placeholder)

    workbook.save(output_file)
    return anonymized_terms

def anonymize_csv(input_file, output_file, sensitive_terms):
    anonymized_terms = {}

    for term in sensitive_terms:
        placeholder = generate_random_string(len(term))
        anonymized_terms[term] = placeholder

    with open(input_file, 'r', encoding='utf-8') as infile, open(output_file, 'w', encoding='utf-8', newline='') as outfile:
        reader = csv.reader(infile)
        writer = csv.writer(outfile)

        for row in reader:
            new_row = []
            for cell in row:
                for term in sensitive_terms:
                    cell = cell.replace(term, anonymized_terms[term])
                new_row.append(cell)
            writer.writerow(new_row)

    return anonymized_terms

def save_mapping_to_file(anonymized_terms, mapping_file):
    with open(mapping_file, 'w', encoding='utf-8') as f:
        for original, anonymized in anonymized_terms.items():
            f.write(f"{original} -> {anonymized}\n")

def select_file(title="Select a file"):
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(title=title)
    return file_path

def load_sensitive_terms(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        terms = [term.strip() for term in f.readlines()]
    return terms

# Prompt the user to select a text file containing the list of sensitive terms
terms_file = select_file("Select a text file containing the list of sensitive terms")
sensitive_terms = load_sensitive_terms(terms_file)

# Prompt the user to select the document to anonymize
input_file = select_file("Select the document to anonymize")
output_folder = "Output Files"
os.makedirs(output_folder, exist_ok=True)
output_file = os.path.join(output_folder, os.path.splitext(os.path.basename(input_file))[0] + '_anonymized' + os.path.splitext(input_file)[1])
mapping_file = os.path.join(output_folder, os.path.splitext(os.path.basename(input_file))[0] + '_anonymized_terms_mapping.txt')

file_extension = os.path.splitext(input_file)[1]

if file_extension == '.docx':
    anonymized_terms = anonymize_docx(input_file, output_file, sensitive_terms)
elif file_extension == '.xlsx':
    anonymized_terms = anonymize_xlsx(input_file, output_file, sensitive_terms)
elif file_extension == '.csv':
    anonymized_terms = anonymize_csv(input_file, output_file, sensitive_terms)
else:
    raise ValueError("Unsupported file format. Only .docx, .xlsx, and .csv files are supported.")

save_mapping_to_file(anonymized_terms, mapping_file)
print("Anonymization process completed. Check the output files.")
